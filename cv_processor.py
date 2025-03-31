import os
import re
import uuid
import logging
from collections import Counter
from docx import Document
import PyPDF2
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.colors import yellow

class CVProcessor:
    """
    Classe pour le traitement avancé des CV en fonction des descriptions de poste.
    """
    
    def __init__(self, upload_folder, download_folder):
        """
        Initialise le processeur de CV avec les dossiers de stockage.
        
        Args:
            upload_folder (str): Chemin vers le dossier de téléchargement des CV
            download_folder (str): Chemin vers le dossier de stockage des CV adaptés
        """
        self.upload_folder = upload_folder
        self.download_folder = download_folder
        
        # Configuration du logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger('CVProcessor')
        
        # Assurer que les dossiers existent
        os.makedirs(upload_folder, exist_ok=True)
        os.makedirs(download_folder, exist_ok=True)
        
        # Mots vides français
        self.stopwords = [
            "le", "la", "les", "un", "une", "des", "et", "ou", "de", "du", "au", "aux",
            "ce", "cette", "ces", "mon", "ton", "son", "notre", "votre", "leur",
            "pour", "par", "sur", "sous", "dans", "avec", "sans", "en", "à", "qui",
            "que", "quoi", "dont", "où", "comment", "pourquoi", "quand", "est", "sont",
            "sera", "seront", "été", "avoir", "être", "faire", "plus", "moins", "très",
            "si", "tout", "tous", "toute", "toutes", "autre", "autres", "même", "aussi",
            "alors", "donc", "car", "mais", "ou", "ni", "ne", "pas", "plus", "moins"
        ]
        
        # Termes importants pour les CV
        self.cv_important_terms = [
            "expérience", "compétence", "formation", "diplôme", "certification",
            "projet", "responsabilité", "gestion", "développement", "analyse",
            "conception", "mise en œuvre", "coordination", "direction", "management",
            "leadership", "stratégie", "objectif", "résultat", "performance"
        ]
    
    def extract_keywords_from_job_description(self, job_description):
        """
        Extrait les mots-clés importants d'une description de poste avec une analyse avancée.
        
        Args:
            job_description (str): Description du poste
            
        Returns:
            dict: Dictionnaire des mots-clés par catégorie avec leur score d'importance
        """
        self.logger.info("Extraction des mots-clés de la description de poste")
        
        # Nettoyage du texte
        job_description = job_description.lower()
        
        # Extraction des mots
        words = re.findall(r'\b\w{3,}\b', job_description)
        
        # Filtrage des mots vides et comptage des occurrences
        filtered_words = [word for word in words if word not in self.stopwords]
        word_counts = Counter(filtered_words)
        
        # Identification des sections importantes
        sections = {
            "compétences": self._extract_section(job_description, ["compétences", "qualifications", "profil", "requis"]),
            "responsabilités": self._extract_section(job_description, ["responsabilités", "missions", "tâches", "rôle"]),
            "formation": self._extract_section(job_description, ["formation", "diplôme", "études", "éducation"]),
            "expérience": self._extract_section(job_description, ["expérience", "parcours", "antécédents"])
        }
        
        # Analyse des mots-clés par section
        keywords = {}
        
        # Mots-clés généraux (apparaissant au moins 2 fois)
        keywords["général"] = {word: count for word, count in word_counts.items() if count >= 2}
        
        # Mots-clés par section
        for section_name, section_text in sections.items():
            if section_text:
                section_words = re.findall(r'\b\w{3,}\b', section_text.lower())
                section_filtered = [word for word in section_words if word not in self.stopwords]
                section_counts = Counter(section_filtered)
                keywords[section_name] = {word: count for word, count in section_counts.items()}
        
        # Mots longs (potentiellement importants)
        long_words = {word: 1 for word in filtered_words if len(word) > 7 and word not in keywords["général"]}
        keywords["termes_spécifiques"] = long_words
        
        # Termes importants pour les CV
        cv_terms = {word: 2 for word in self.cv_important_terms if word in filtered_words}
        keywords["termes_cv"] = cv_terms
        
        self.logger.info(f"Extraction terminée: {sum(len(v) for v in keywords.values())} mots-clés trouvés")
        return keywords
    
    def _extract_section(self, text, section_markers):
        """
        Extrait une section spécifique du texte basée sur des marqueurs.
        
        Args:
            text (str): Texte complet
            section_markers (list): Liste des marqueurs de section
            
        Returns:
            str: Texte de la section ou chaîne vide si non trouvée
        """
        for marker in section_markers:
            pattern = rf'{marker}\s*:(.+?)(?:\n\n|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
        return ""
    
    def _annotate_pdf_page(self, page, text, keywords):
        """
        Ajoute des annotations sur une page PDF pour mettre en évidence les mots-clés.
        
        Args:
            page: Page PDF à annoter
            text (str): Texte extrait de la page
            keywords (list): Liste des mots-clés à mettre en évidence
        """
        # Créer un canvas pour les annotations
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=page.mediabox)
        
        # Trouver les positions des mots-clés dans le texte
        for keyword in keywords:
            start = 0
            while True:
                pos = text.lower().find(keyword.lower(), start)
                if pos == -1:
                    break
                    
                # Calculer la position approximative sur la page
                # Note: Cette méthode est simplifiée et pourrait être améliorée
                x = pos * 7  # Approximation grossière
                y = page.mediabox.height - (pos // 50 * 12)  # Approximation grossière
                
                # Ajouter une annotation en surbrillance
                can.setFillColor(yellow, alpha=0.3)
                can.rect(x, y, len(keyword) * 7, 12, fill=True)
                
                start = pos + 1
        
        can.save()
        packet.seek(0)
        
        # Créer une nouvelle page avec les annotations
        new_pdf = PyPDF2.PdfReader(packet)
        page.merge_page(new_pdf.pages[0])
        
        return page

    def adapt_cv(self, cv_path, job_description):
        """
        Adapte un CV en fonction d'une description de poste avec mise en évidence avancée.
        
        Args:
            cv_path (str): Chemin vers le fichier CV
            job_description (str): Description du poste
            
        Returns:
            dict: Informations sur le CV adapté (nom du fichier, statistiques)
        """
        self.logger.info(f"Adaptation du CV: {cv_path}")
        
        # Vérifier que le fichier existe
        if not os.path.exists(cv_path):
            error_msg = f"Le fichier CV n'existe pas: {cv_path}"
            self.logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        # Extraire les mots-clés de la description du poste
        keywords_dict = self.extract_keywords_from_job_description(job_description)
        
        # Créer une liste plate de tous les mots-clés
        all_keywords = []
        for category, words in keywords_dict.items():
            all_keywords.extend(words.keys())
        
        # Éliminer les doublons
        all_keywords = list(set(all_keywords))
        
        # Détecter le type de fichier
        file_ext = os.path.splitext(cv_path)[1].lower()
        
        try:
            # Détecter et traiter selon le type de fichier
            if file_ext == '.docx':
                # Traitement des fichiers Word
                doc = Document(cv_path)
                
                # Statistiques
                stats = {
                    "total_keywords": len(all_keywords),
                    "highlighted_keywords": 0,
                    "paragraphs_modified": 0
                }
                
                # Adapter le CV en mettant en évidence les compétences pertinentes
                for paragraph in doc.paragraphs:
                    paragraph_modified = False
                    
                    for keyword in all_keywords:
                        if keyword in paragraph.text.lower():
                            # Mettre en gras les mots-clés trouvés
                            pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                            
                            # Sauvegarde du texte original
                            original_text = paragraph.text
                            
                            # Effacer le contenu du paragraphe
                            paragraph.clear()
                            
                            # Trouver toutes les occurrences du mot-clé
                            matches = list(pattern.finditer(original_text))
                            
                            # Si aucune correspondance, continuer
                            if not matches:
                                paragraph.add_run(original_text)
                                continue
                            
                            # Ajouter le texte avec les mots-clés en gras
                            last_end = 0
                            for match in matches:
                                # Ajouter le texte avant le mot-clé
                                if match.start() > last_end:
                                    paragraph.add_run(original_text[last_end:match.start()])
                                
                                # Ajouter le mot-clé souligné
                                paragraph.add_run(original_text[match.start():match.end()]).underline = True
                                stats["highlighted_keywords"] += 1
                                
                                last_end = match.end()
                            
                            # Ajouter le reste du texte
                            if last_end < len(original_text):
                                paragraph.add_run(original_text[last_end:])
                            
                            paragraph_modified = True
                    
                    if paragraph_modified:
                        stats["paragraphs_modified"] += 1
                
            elif file_ext == '.pdf':
                # Traitement des fichiers PDF
                with open(cv_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    pdf_writer = PyPDF2.PdfWriter()
                    
                    # Statistiques
                    stats = {
                        "total_keywords": len(all_keywords),
                        "highlighted_keywords": 0,
                        "pages_modified": 0
                    }
                    
                    # Traiter chaque page du PDF
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        
                        if text:
                            # Vérifier les mots-clés dans le texte
                            for keyword in all_keywords:
                                if keyword in text.lower():
                                    stats["highlighted_keywords"] += text.lower().count(keyword)
                            
                            if stats["highlighted_keywords"] > 0:
                                # Annoter la page avec les mots-clés trouvés
                                page = self._annotate_pdf_page(page, text, all_keywords)
                                stats["pages_modified"] += 1
                            
                            # Ajouter la page au writer
                            pdf_writer.add_page(page)
            
            # Générer un nom de fichier unique pour le CV adapté
            output_filename = f"cv_adapte_{uuid.uuid4().hex}{file_ext}"
            output_path = os.path.join(self.download_folder, output_filename)
            
            # Assurer que le dossier de destination existe
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Sauvegarder le document modifié selon le type de fichier
            if file_ext == '.docx':
                doc.save(output_path)
            elif file_ext == '.pdf':
                # Pour les PDF, on sauvegarde le fichier modifié
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
            
            # Vérifier que le fichier a bien été créé
            if not os.path.exists(output_path):
                error_msg = f"Échec de la création du fichier adapté: {output_path}"
                self.logger.error(error_msg)
                raise IOError(error_msg)
            
            self.logger.info(f"CV adapté avec succès: {output_filename}")
            self.logger.info(f"Statistiques: {stats}")
            
            return {
                "filename": output_filename,
                "stats": stats,
                "keywords": keywords_dict
            }
            
        except Exception as e:
            error_msg = f"Erreur lors de l'adaptation du CV: {str(e)}"
            self.logger.error(error_msg)
            raise Exception(error_msg)
    
    def analyze_job_description(self, job_description):
        """
        Analyse une description de poste pour extraire des informations structurées.
        
        Args:
            job_description (str): Description du poste
            
        Returns:
            dict: Informations extraites de la description du poste
        """
        self.logger.info("Analyse de la description de poste")
        
        # Extraction des compétences techniques potentielles
        tech_skills_pattern = r'compétences\s+techniques.*?:(.*?)(?:\n\n|\Z)'
        tech_skills_match = re.search(tech_skills_pattern, job_description, re.IGNORECASE | re.DOTALL)
        tech_skills = []
        if tech_skills_match:
            tech_skills_text = tech_skills_match.group(1)
            tech_skills = [skill.strip() for skill in re.split(r'[,;•]', tech_skills_text) if skill.strip()]
        
        # Extraction du niveau d'expérience requis
        experience_pattern = r'(\d+)[+]?\s+ans?\s+d\'expérience'
        experience_match = re.search(experience_pattern, job_description, re.IGNORECASE)
        experience_years = experience_match.group(1) if experience_match else None
        
        # Extraction du niveau d'études requis
        education_patterns = [
            r'(bac\s*[+]\s*\d+)',
            r'(master|licence|doctorat)',
            r'(diplôme\s+d\'ingénieur)'
        ]
        education = None
        for pattern in education_patterns:
            match = re.search(pattern, job_description, re.IGNORECASE)
            if match:
                education = match.group(1)
                break
        
        # Extraction du type de contrat
        contract_pattern = r'\b(cdi|cdd|stage|alternance|freelance|intérim)\b'
        contract_match = re.search(contract_pattern, job_description, re.IGNORECASE)
        contract_type = contract_match.group(1).upper() if contract_match else None
        
        # Résultats de l'analyse
        analysis = {
            "keywords": self.extract_keywords_from_job_description(job_description),
            "technical_skills": tech_skills,
            "experience_years": experience_years,
            "education": education,
            "contract_type": contract_type
        }
        
        self.logger.info("Analyse de la description de poste terminée")
        return analysis
